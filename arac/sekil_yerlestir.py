#!/usr/bin/env python3
"""
Üretilen PNG şekilleri, kitabın .docx bölümlerinde ilgili kod bloğunun altına
"Şekil N.k: …" açıklamasıyla yerleştirir.

Kullanım
--------
    python kod/arac/sekil_yerlestir.py --bolumler ./bolumler
    python kod/arac/sekil_yerlestir.py --bolumler ./bolumler --sekiller ./sekiller
    python kod/arac/sekil_yerlestir.py --bolumler ./bolumler --onizleme

Gerekli: pip install python-docx

Şeklin yeri, .docx içindeki  "📄 Tam kod: kod/bolumNN/dosya.py"  satırından bulunur.
Her çalıştırmada şekil numaraları 1'den başlar; bu yüzden betik, bölümde daha önce
yerleştirilmiş şekilleri varsayılan olarak temizler (--koru ile kapatılabilir).
"""
import argparse
import re
import shutil
import sys
from collections import defaultdict
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit('HATA: python-docx kurulu değil.  ->  pip install python-docx')

BURASI = Path(__file__).resolve().parent

ap = argparse.ArgumentParser()
ap.add_argument('--bolumler', type=Path, default=None,
                help='.docx bölüm dosyalarının bulunduğu klasör (ZORUNLU)')
ap.add_argument('--sekiller', type=Path, default=None, help='PNG klasörü')
ap.add_argument('--onizleme', action='store_true', help='dosyaları değiştirme')
ap.add_argument('--koru', action='store_true',
                help='daha önce eklenmiş şekilleri silme (numaralar çakışabilir)')
A = ap.parse_args()

SEK = (A.sekiller or BURASI.parent.parent / 'sekiller').resolve()
BOL = A.bolumler.resolve() if A.bolumler else None

if BOL is None:
    for aday in (Path.cwd() / 'bolumler', BURASI.parent.parent / 'bolumler'):
        if aday.is_dir() and list(aday.glob('*.docx')):
            BOL = aday.resolve()
            break
if BOL is None or not BOL.is_dir():
    sys.exit('HATA: bölüm klasörü bulunamadı.\n'
             '      --bolumler ile verin, örn:  --bolumler ./bolumler')

docxler = sorted(BOL.glob('*.docx'))
if not docxler:
    sys.exit(f'HATA: {BOL} içinde .docx yok.')
if not SEK.is_dir() or not list(SEK.glob('*.png')):
    sys.exit(f'HATA: {SEK} içinde PNG yok. Önce sekil_uret.py çalıştırın.')

print(f'bölümler : {BOL}  ({len(docxler)} dosya)')
print(f'şekiller : {SEK}')

# PNG'leri kaynak kod dosyasına göre grupla
gruplar = defaultdict(list)
for p in sorted(SEK.glob('*.png')):
    kok = p.stem.rsplit('__', 1)[0]
    if '__' not in kok:
        continue
    bl, rel = kok.split('__', 1)
    gruplar[f'{bl}/{rel}.py'].append(p)
print(f'{len(gruplar)} kod dosyası için {sum(len(v) for v in gruplar.values())} şekil\n')

REF = re.compile(r'Tam kod.*?:\s*kod/(bolum\d+/\S+)')
CAP = re.compile(r'^\s*Şekil\s+\d+\.\d+\s*:')


def temiz(t):
    t = re.sub(r'^\d+(\.\d+)*\.?\s*', '', t.strip())
    t = re.sub(r'^[A-Z]\.\s*', '', t)
    t = re.sub(r'^Python Uygulaması\s*[—–\-:]\s*', '', t)
    return t.strip(' —–-:') or 'Uygulama çıktısı'


def aciklama_p(metin):
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '40')
    sp.set(qn('w:after'), '200')
    ppr.append(sp)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    ppr.append(jc)
    p.append(ppr)
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rpr.append(OxmlElement('w:i'))
    for tag, val in (('w:sz', '17'), ('w:color', '444444')):
        e = OxmlElement(tag)
        e.set(qn('w:val'), val)
        rpr.append(e)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = metin
    r.append(t)
    p.append(r)
    return p


toplam = 0
for f in docxler:
    m = re.search(r'(\d+)', f.name)
    if not m:
        continue
    n = int(m.group(1))
    d = Document(str(f))

    if not A.koru and not A.onizleme:
        silinecek = []
        for p in d.paragraphs:
            if CAP.match(p.text or ''):
                silinecek.append(p._p)
                onc = p._p.getprevious()
                if onc is not None and onc.findall('.//' + qn('w:drawing')):
                    silinecek.append(onc)
        for e in silinecek:
            if e.getparent() is not None:
                e.getparent().remove(e)

    hedefler, son = [], ''
    for p in d.paragraphs:
        if (p.style.name or '').startswith('Heading'):
            son = temiz(p.text)
        mm = REF.search(p.text or '')
        if mm and mm.group(1) in gruplar:
            hedefler.append((p, mm.group(1), son))
    if not hedefler:
        continue

    if not A.onizleme:
        shutil.copy(f, f.with_suffix('.docx.yedek'))

    no = 0
    for p, rel, baslik in hedefler:
        capa = p._p
        for png in gruplar[rel]:
            no += 1
            toplam += 1
            if A.onizleme:
                continue
            gp = d.add_paragraph()
            gp.alignment = 1
            gp.add_run().add_picture(str(png), width=Inches(6.1))
            capa.addnext(gp._p)
            cap = aciklama_p(f'Şekil {n}.{no}: {baslik}')
            gp._p.addnext(cap)
            capa = cap
    if not A.onizleme:
        d.save(str(f))
    print(f'  {f.name[:56]:<58} {no} şekil')

print(f'\nTOPLAM {toplam} şekil' + (' (ÖNİZLEME — dosyalar değişmedi)' if A.onizleme
                                    else ' yerleştirildi. Yedekler: *.docx.yedek'))
